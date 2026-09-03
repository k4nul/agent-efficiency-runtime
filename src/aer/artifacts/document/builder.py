from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

from aer.errors import AerError
from aer.limits import MAX_ARTIFACT_ELEMENTS, MAX_IMAGE_PIXELS
from aer.paths import atomic_write_bytes, ensure_regular_input


def _bookmark(paragraphs: list[Any], block_id: str, bookmark_id: int) -> None:
    """Wrap one logical block, which may span multiple paragraphs, in a bookmark."""

    name = f"aer_{block_id.replace('-', '_')}"
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraphs[0]._p.insert(0, start)
    paragraphs[-1]._p.append(end)


def _table_id(table: Any, block_id: str) -> None:
    properties = table._tbl.tblPr
    caption = OxmlElement("w:tblCaption")
    caption.set(qn("w:val"), f"aer:{block_id}")
    properties.append(caption)


def _page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _paragraph(document: Any, text: str, *, style: str | None = None) -> Any:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Noto Sans CJK KR"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK KR")
    return paragraph


def _set_cell_text(cell: Any, value: Any, *, header: bool = False) -> None:
    cell.text = str(value)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Noto Sans CJK KR"
            run.font.size = Pt(9)
            run.font.bold = header
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK KR")
    if header:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "193A5A")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)


def _validate_document_limits(content: list[Any]) -> None:
    rendered_elements = len(content)
    if rendered_elements > MAX_ARTIFACT_ELEMENTS:
        raise AerError(
            "LIMIT_EXCEEDED",
            "Document exceeds the block-count limit.",
            "document.build",
            "/content",
            {"blocks": len(content), "limit": MAX_ARTIFACT_ELEMENTS},
        )
    for index, block in enumerate(content):
        if not isinstance(block, dict):
            continue
        kind = str(block.get("type", block.get("block", "paragraph")))
        if kind in {"bullets", "numbered-list", "source-list"}:
            items = block.get("items", [])
            if not isinstance(items, list):
                raise AerError(
                    "INVALID_SPEC",
                    "Document list items must be an array.",
                    "document.build",
                    f"/content/{index}/items",
                )
            rendered_elements += max(1, len(items)) - 1
        elif kind == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if not isinstance(headers, list) or not isinstance(rows, list):
                raise AerError(
                    "INVALID_SPEC",
                    "Document table headers and rows must be arrays.",
                    "document.build",
                    f"/content/{index}",
                )
            column_count = len(headers)
            for row_index, row in enumerate(rows):
                if not isinstance(row, list):
                    raise AerError(
                        "INVALID_SPEC",
                        "Document table rows must be arrays.",
                        "document.build",
                        f"/content/{index}/rows/{row_index}",
                    )
                column_count = max(column_count, len(row))
            table_cells = (len(rows) + (1 if headers else 0)) * column_count
            rendered_elements += max(0, table_cells - 1)
        if rendered_elements > MAX_ARTIFACT_ELEMENTS:
            raise AerError(
                "LIMIT_EXCEEDED",
                "Document exceeds the rendered-element limit.",
                "document.build",
                f"/content/{index}",
                {"elements": rendered_elements, "limit": MAX_ARTIFACT_ELEMENTS},
            )


def build_document(
    spec: dict[str, Any], output: Path, *, spec_dir: Path
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    content = spec.get("content")
    if not isinstance(content, list) or not content:
        raise AerError(
            "INVALID_SPEC",
            "Document content must be a non-empty array.",
            "document.build",
            "/content",
        )
    _validate_document_limits(content)
    document = Document()
    metadata = spec.get("metadata", {})
    document.core_properties.title = str(metadata.get("title", ""))
    document.core_properties.subject = "Built by Agent Efficiency Runtime"
    page = spec.get("page", {})
    for section in document.sections:
        if str(page.get("size", "A4")).upper() == "LETTER":
            section.page_width, section.page_height = Inches(8.5), Inches(11)
        else:
            section.page_width, section.page_height = Cm(21), Cm(29.7)
        margins = page.get("margins", {})
        section.top_margin = Cm(float(margins.get("top", 2.2)))
        section.bottom_margin = Cm(float(margins.get("bottom", 2.2)))
        section.left_margin = Cm(float(margins.get("left", 2.2)))
        section.right_margin = Cm(float(margins.get("right", 2.2)))
    styles = document.styles
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Noto Sans CJK KR"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK KR")
    if "AER Callout" not in styles:
        callout = styles.add_style("AER Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Noto Sans CJK KR"
        callout.font.color.rgb = RGBColor(25, 42, 67)
    header = spec.get("header")
    footer = spec.get("footer")
    for section in document.sections:
        if header:
            section.header.paragraphs[0].text = str(header)
        if footer:
            section.footer.paragraphs[0].text = str(footer)
        if spec.get("page_numbers", True):
            number_paragraph = section.footer.add_paragraph()
            _page_number(number_paragraph)
    elements: list[dict[str, Any]] = []
    seen_ids: set[str] = set()
    seen_selector_keys: set[str] = set()
    bookmark_id = 1
    for index, block in enumerate(content):
        if not isinstance(block, dict):
            raise AerError(
                "INVALID_SPEC",
                "Each document block must be an object.",
                "document.build",
                f"/content/{index}",
            )
        block_id = str(block.get("id", f"block-{index + 1}"))
        selector_key = block_id.replace("-", "_")
        if (
            not block_id
            or block_id in seen_ids
            or selector_key in seen_selector_keys
            or "/" in block_id
        ):
            raise AerError(
                "INVALID_SPEC",
                "Block IDs must be unique after stable-selector normalization and cannot contain '/'.",
                "document.build",
                f"/content/{index}/id",
            )
        seen_ids.add(block_id)
        seen_selector_keys.add(selector_key)
        kind = str(block.get("type", block.get("block", "paragraph")))
        element: Any | None = None
        block_paragraphs: list[Any] = []
        if kind == "title":
            element = _paragraph(
                document, str(block.get("text", block.get("title", ""))), style="Title"
            )
        elif kind == "heading":
            level = max(1, min(9, int(block.get("level", 1))))
            element = _paragraph(
                document, str(block.get("text", "")), style=f"Heading {min(level, 3)}"
            )
            element.paragraph_format.keep_with_next = True
        elif kind == "paragraph":
            element = _paragraph(document, str(block.get("text", "")))
        elif kind in {"bullets", "numbered-list"}:
            items = list(block.get("items", []))
            for item_index, item in enumerate(items):
                paragraph = _paragraph(
                    document,
                    str(item.get("text", "")) if isinstance(item, dict) else str(item),
                    style="List Bullet" if kind == "bullets" else "List Number",
                )
                block_paragraphs.append(paragraph)
                if item_index == 0:
                    element = paragraph
            if element is None:
                element = _paragraph(
                    document, "", style="List Bullet" if kind == "bullets" else "List Number"
                )
                block_paragraphs.append(element)
        elif kind == "table":
            headers = list(block.get("headers", []))
            rows = list(block.get("rows", []))
            column_count = len(headers) or (len(rows[0]) if rows else 0)
            if column_count == 0:
                raise AerError(
                    "INVALID_SPEC",
                    "Table requires headers or rows.",
                    "document.build",
                    f"/content/{index}",
                )
            table = document.add_table(rows=len(rows) + (1 if headers else 0), cols=column_count)
            table.style = str(block.get("style", "Table Grid"))
            cursor = 0
            if headers:
                for column, value in enumerate(headers):
                    _set_cell_text(table.cell(0, column), value, header=True)
                cursor = 1
            for row_index, row in enumerate(rows, start=cursor):
                for column in range(column_count):
                    _set_cell_text(
                        table.cell(row_index, column), row[column] if column < len(row) else ""
                    )
            _table_id(table, block_id)
            elements.append(
                {
                    "id": block_id,
                    "type": "table",
                    "selector": f"block:id={block_id}",
                    "index": index,
                }
            )
            continue
        elif kind == "image":
            raw_source = block.get("source")
            if not raw_source:
                raise AerError(
                    "INVALID_SPEC",
                    "Image block requires source.",
                    "document.build",
                    f"/content/{index}/source",
                )
            source = (
                spec_dir / str(raw_source)
                if not Path(str(raw_source)).is_absolute()
                else Path(str(raw_source))
            )
            source = ensure_regular_input(source, operation="document.build")
            try:
                with Image.open(source) as image:
                    if image.width * image.height > MAX_IMAGE_PIXELS:
                        raise AerError(
                            "LIMIT_EXCEEDED",
                            "Image pixel count exceeds the safety limit.",
                            "document.build",
                            str(source),
                        )
                    image.verify()
            except AerError:
                raise
            except Exception as exc:
                raise AerError(
                    "CORRUPT_FILE",
                    f"Cannot open image: {exc}",
                    "document.build",
                    str(source),
                ) from exc
            element = document.add_paragraph()
            element.alignment = WD_ALIGN_PARAGRAPH.CENTER
            element.add_run().add_picture(str(source), width=Inches(float(block.get("width", 5.5))))
        elif kind == "caption":
            element = _paragraph(document, str(block.get("text", "")), style="Caption")
            element.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "quote":
            element = _paragraph(
                document, str(block.get("text", block.get("quote", ""))), style="Quote"
            )
        elif kind == "callout":
            element = _paragraph(document, str(block.get("text", "")), style="AER Callout")
            properties = element._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "E9F0F7")
            properties.append(shading)
        elif kind == "source-list":
            for item_index, item in enumerate(block.get("items", [])):
                paragraph = _paragraph(document, str(item), style="List Number")
                block_paragraphs.append(paragraph)
                if item_index == 0:
                    element = paragraph
            if element is None:
                element = _paragraph(document, "")
                block_paragraphs.append(element)
        elif kind == "page-break":
            element = document.add_paragraph()
            element.add_run().add_break(WD_BREAK.PAGE)
        elif kind == "section-break":
            document.add_section(WD_SECTION.NEW_PAGE)
            element = document.add_paragraph()
        else:
            raise AerError(
                "INVALID_SPEC",
                f"Unsupported document block: {kind}",
                "document.build",
                f"/content/{index}/type",
            )
        if not block_paragraphs:
            block_paragraphs.append(element)
        _bookmark(block_paragraphs, block_id, bookmark_id)
        bookmark_id += 1
        elements.append(
            {"id": block_id, "type": kind, "selector": f"block:id={block_id}", "index": index}
        )
    buffer = io.BytesIO()
    document.save(buffer)
    atomic_write_bytes(output, buffer.getvalue())
    return elements, []
